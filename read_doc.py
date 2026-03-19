#!/usr/bin/env python3
from docx import Document

doc = Document('test-reports/Assignment2_Test_Report.docx')

print('='*80)
print('DOCUMENT STRUCTURE')
print('='*80)

for i, para in enumerate(doc.paragraphs[:200]):
    text = para.text.strip()
    if not text:
        continue
    
    if para.style.name.startswith('Heading'):
        level = para.style.name.replace('Heading ', '')
        indent = '  ' * (int(level) - 1 if level.isdigit() else 0)
        print(f'\n{indent}[H{level}] {text}')
    elif '📷' in text or 'SCREENSHOT' in text or '[INSERT' in text.upper():
        print(f'    📸 {text[:80]}')
    elif len(text) > 30:
        print(f'    {text[:70]}...')

print(f'\n\nTotal paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')

# Count images
image_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run._element.xpath('.//pic:pic'):
            image_count += 1
            
print(f'Total images: {image_count}')
